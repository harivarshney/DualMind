"""
Result Exporter Service
Handles exporting results to various formats (PDF, text, etc.)
"""

import os
from datetime import datetime
from typing import Optional

class ResultExporter:
    def __init__(self):
        """Initialize the result exporter"""
        self.default_font_size = 12
        self.default_margin = 72  # 1 inch margins
        
    def save_as_pdf(self, content: str, output_path: str) -> None:
        """
        Save text content as a formatted PDF file
        
        Args:
            content (str): Text content to save
            output_path (str): Path where PDF should be saved
        """
        
        if not content.strip():
            raise ValueError("Cannot save empty content")
            
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors
            
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=self.default_margin,
                leftMargin=self.default_margin,
                topMargin=self.default_margin,
                bottomMargin=self.default_margin
            )
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1,  # Center alignment
                textColor=colors.darkblue
            )
            
            content_style = ParagraphStyle(
                'CustomContent',
                parent=styles['Normal'],
                fontSize=self.default_font_size,
                spaceAfter=12,
                leading=14
            )
            
            # Build PDF content
            story = []
            
            # Add title
            story.append(Paragraph("DualMind - Exported Results", title_style))
            story.append(Spacer(1, 20))
            
            # Add metadata
            metadata = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            story.append(Paragraph(metadata, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Add content
            # Split content by paragraphs and add each as a separate paragraph
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Escape special characters for reportlab
                    escaped_para = self._escape_for_pdf(para.strip())
                    story.append(Paragraph(escaped_para, content_style))
                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
        except ImportError:
            raise ImportError(
                "reportlab library is required for PDF export. "
                "Install it with: pip install reportlab"
            )
        except Exception as e:
            raise Exception(f"Failed to create PDF: {str(e)}")
        
    def save_as_text(self, content: str, output_path: str) -> None:
        """
        Save content as a plain text file
        
        Args:
            content (str): Text content to save
            output_path (str): Path where text file should be saved
        """
        if not content.strip():
            raise ValueError("Cannot save empty content")
            
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("DUALMIND - EXPORTED RESULTS\n")
                f.write("=" * 50 + "\n\n")
                f.write(f"Exported: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"File: {os.path.basename(output_path)}\n\n")
                f.write("CONTENT:\n")
                f.write("-" * 20 + "\n")
                f.write(content)
                f.write("\n\n" + "=" * 50)
                
        except Exception as e:
            raise Exception(f"Failed to save text file: {str(e)}")
            
    def _escape_for_pdf(self, text: str) -> str:
        """
        Escape special characters for PDF generation
        
        Args:
            text (str): Raw text content
            
        Returns:
            str: Escaped text safe for PDF
        """
        # Escape special characters that might cause issues in reportlab
        replacements = {
            '&': '&amp;',
            '<': '&lt;',
            '>': '&gt;',
            '"': '&quot;',
            "'": '&apos;'
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
            
        return text
        
    def save_as_word(self, content: str, output_path: str) -> None:
        """
        Save text content as a formatted Word document
        
        Args:
            content (str): Text content to save
            output_path (str): Path where Word document should be saved
        """
        
        if not content.strip():
            raise ValueError("Cannot save empty content")
            
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement, qn
            
            # Create Word document
            doc = Document()
            
            # Set document margins (1 inch)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add title
            title = doc.add_heading('DualMind - Exported Results', level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata paragraph
            metadata_para = doc.add_paragraph()
            metadata_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            metadata_run = metadata_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            metadata_run.italic = True
            
            # Add some space
            doc.add_paragraph()
            
            # Add content
            # Split content by paragraphs and add each as a separate paragraph
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = doc.add_paragraph()
                    
                    # Check if it's a heading (starts with emojis or special characters)
                    if any(char in para_text[:10] for char in ['🎥', '📽️', '🔗', '👤', '⏱️', '🌍', '🕒', '📝', '=']):
                        # Format as heading or special paragraph
                        run = para.add_run(para_text.strip())
                        if '=' in para_text:
                            run.bold = True
                            run.font.size = Pt(14)
                        else:
                            run.font.size = Pt(12)
                    else:
                        # Regular content
                        run = para.add_run(para_text.strip())
                        run.font.size = Pt(11)
                    
                    # Set line spacing
                    para.paragraph_format.line_spacing = 1.15
                    para.paragraph_format.space_after = Pt(6)
            
            # Save document
            doc.save(output_path)
            
        except ImportError:
            raise ImportError(
                "python-docx library is required for Word export. "
                "Install it with: pip install python-docx"
            )
        except Exception as e:
            raise Exception(f"Failed to create Word document: {str(e)}")

    def get_supported_formats(self) -> list:
        """
        Get list of supported export formats
        
        Returns:
            list: List of supported format extensions
        """
        return ['.pdf', '.docx', '.txt']
        
    def validate_output_path(self, output_path: str) -> bool:
        """
        Validate if the output path is valid and writable
        
        Args:
            output_path (str): Path to validate
            
        Returns:
            bool: True if path is valid, False otherwise
        """
        try:
            # Check if directory exists or can be created
            directory = os.path.dirname(output_path)
            if directory and not os.path.exists(directory):
                try:
                    os.makedirs(directory, exist_ok=True)
                except:
                    return False
                    
            # Check if we can write to the location
            test_file = output_path + '.test'
            try:
                with open(test_file, 'w') as f:
                    f.write('test')
                os.remove(test_file)
                return True
            except:
                return False
                
        except:
            return False
            
    def create_export_filename(self, base_name: str, format_ext: str) -> str:
        """
        Create a standardized export filename
        
        Args:
            base_name (str): Base name for the file
            format_ext (str): File format extension (.pdf, .txt)
            
        Returns:
            str: Formatted filename with timestamp
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_base = "".join(c for c in base_name if c.isalnum() or c in (' ', '-', '_'))
        clean_base = clean_base.replace(' ', '_')
        
        return f"DualMind_{clean_base}_{timestamp}{format_ext}"
        
    def get_file_size_mb(self, file_path: str) -> float:
        """
        Get file size in megabytes
        
        Args:
            file_path (str): Path to file
            
        Returns:
            float: File size in MB
        """
        try:
            if os.path.exists(file_path):
                size_bytes = os.path.getsize(file_path)
                return round(size_bytes / (1024 * 1024), 2)
            return 0.0
        except:
            return 0.0
